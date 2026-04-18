"""Populate the RL project template with Week 1 content."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

doc = Document("deliverables/Reinforcement_Learning_Project_Template.docx")

# --- Fill in cover page info ---
for i, p in enumerate(doc.paragraphs):
    if p.text == "Project or System Name":
        p.clear()
        run = p.add_run("Dyna-Q Reinforcement Learning for Magic: The Gathering Commander Deck Construction")
        run.font.size = Pt(14)
        run.bold = True
    elif p.text == "Student name and email":
        p.clear()
        run = p.add_run("Keith Coughlin — krc5765@psu.edu")
        run.font.size = Pt(12)
    elif p.text == "Month, Year":
        p.clear()
        run = p.add_run("March, 2026")
        run.font.size = Pt(12)

# --- Fill in Document Control table ---
table0 = doc.tables[0]
table0.rows[1].cells[0].text = "Keith Coughlin"
table0.rows[1].cells[1].text = "krc5765@psu.edu"
table0.rows[1].cells[2].text = "Author"

# Revision table
table1 = doc.tables[1]
table1.rows[1].cells[0].text = "1.0"
table1.rows[1].cells[1].text = "March 16, 2026"
table1.rows[1].cells[2].text = "Project Assignment 1 — Research goal and RL problem formulation"


def add_heading_after(paragraph, text, level=2):
    """Add a heading paragraph after the given paragraph."""
    new_p = doc.add_paragraph()
    new_p.style = doc.styles[f'Heading {level}']
    new_p.text = text
    # Move it after the target paragraph
    paragraph._element.addnext(new_p._element)
    return new_p


def add_paragraph_after(paragraph, text, style='Normal', bold=False):
    """Add a paragraph after the given paragraph."""
    new_p = doc.add_paragraph()
    new_p.style = doc.styles[style]
    if bold:
        run = new_p.add_run(text)
        run.bold = True
    else:
        new_p.text = text
    paragraph._element.addnext(new_p._element)
    return new_p


# Find "Project Assignment 1" and the note about OpenAI Gym (end of section 1 content area)
pa1_note = None
pa2_start = None
for i, p in enumerate(doc.paragraphs):
    if "Students are welcome to choose" in p.text:
        pa1_note = p
    if p.text.strip() == "Example (Default project): Cliff Walking grid-world problem discussed in lesson 7. Refer to figure 7.1":
        pa1_note = p  # Use the last paragraph before Assignment 2
    if p.text.strip() == "Project Assignment 2":
        pa2_start = p
        break

# We'll insert content BEFORE "Project Assignment 2" — i.e., after the example note
# Build content in reverse order since addnext inserts right after the anchor

content_blocks = [
    ("heading2", "1. Research / Business Goal"),
    ("normal", ""),
    ("bold", "Research Question"),
    ("normal", "Can a model-based reinforcement learning agent learn to construct competitive Magic: The Gathering Commander decks from a constrained physical card collection, using only card category abstractions and coarse state representations?"),
    ("normal", ""),
    ("bold", "Motivation"),
    ("normal", "Magic: The Gathering Commander is a 100-card singleton format where deck construction requires balancing dozens of competing concerns: mana curve, removal density, card advantage, synergy with the commander, and win condition coherence. Expert players rely on years of pattern recognition to navigate these trade-offs. Most Commander players own large collections (3,000\u20138,000+ cards) and struggle to build optimized decks from what they already own."),
    ("normal", ""),
    ("normal", "This project applies Dyna-Q reinforcement learning to discover effective deck construction policies. The agent builds decks from a specific physical collection \u2014 not the entire card pool \u2014 providing personalized, actionable deck recommendations grounded in cards the player actually has."),
    ("normal", ""),
    ("bold", "Algorithm Choice: Dyna-Q"),
    ("normal", "Dyna-Q (Sutton, 1991) combines model-free Q-learning with model-based planning. Two properties of this problem make Dyna-Q a strong fit:"),
    ("normal", "1. Simulated experience is cheap. Once card metadata is cached locally, the agent can simulate deck-building episodes with zero external API cost. Dyna-Q\u2019s planning steps exploit this by replaying past transitions through a learned model, dramatically improving sample efficiency over pure Q-learning."),
    ("normal", "2. The state-action space is tractable but not trivial. With approximately 729 discretized states and 50 card category actions, the Q-table has roughly 36,450 entries \u2014 large enough that pure random exploration would be slow, but small enough that tabular methods converge reliably. Dyna-Q\u2019s model-based planning fills exactly this gap."),
    ("normal", ""),
    ("heading2", "2. RL Problem Formulation"),
    ("normal", ""),
    ("heading3", "2.1 Agent"),
    ("normal", "The agent is a Dyna-Q learner that builds a Commander deck one card category at a time. At each time step, it observes the current state of the partially-built deck and selects a card category to add. A downstream heuristic then picks the best specific card from the player\u2019s collection for that category."),
    ("normal", ""),
    ("normal", "The agent maintains three components:"),
    ("normal", "\u2022 A Q-table Q(s, a) mapping state-action pairs to expected cumulative reward"),
    ("normal", "\u2022 A learned transition model Model(s, a) \u2192 (s\u2019, r) used for Dyna planning steps"),
    ("normal", "\u2022 An \u03b5-greedy exploration policy that decays over training"),
    ("normal", ""),
    ("heading3", "2.2 Environment"),
    ("normal", "The environment represents the deck-building process for a single commander, implemented using the Gymnasium library\u2019s gymnasium.Env interface (Towers et al., 2024)."),
    ("normal", ""),
    ("bold", "Episode structure:"),
    ("normal", "1. A commander is selected from a curated roster of 10\u201315 commanders covering distinct color identities and archetypes."),
    ("normal", "2. The agent fills all 99 deck slots one at a time (1 commander + 99 selected cards = 100 total). Every slot is a strategic decision, including lands \u2014 the agent chooses among non-basic land categories (shock lands, fetch lands, dual lands, utility lands, etc.) alongside spells and creatures."),
    ("normal", "3. The episode terminates when all 99 slots are filled."),
    ("normal", "4. A terminal reward is computed by submitting the completed deck to an external deck analysis API."),
    ("normal", ""),
    ("bold", "Hard constraints enforced by the environment:"),
    ("normal", "\u2022 Singleton rule: No duplicate card names may appear in the deck (except basic lands)."),
    ("normal", "\u2022 Color identity: Every non-land card must fall within the commander\u2019s color identity."),
    ("normal", "\u2022 Collection-bound: Only cards the player actually owns are available for selection."),
    ("normal", "\u2022 Action masking: If a category has no remaining eligible cards, that action is unavailable to the agent."),
    ("normal", ""),
    ("heading3", "2.3 State Space"),
    ("normal", "The state represents the composition of the partially-built deck using coarsely discretized bins. Rather than tracking individual cards (which would produce an astronomically large state space), the state captures deck-level statistics that are relevant to construction decisions."),
    ("normal", ""),
    ("bold", "Table 1: State Variables"),
    ("normal", "Each variable is discretized into 3 bins (low / medium / high):"),
    ("normal", ""),
    ("normal", "\u2022 Mana curve position (average CMC of selected cards) \u2014 Keeps the deck castable; too high means clunky starts, too low means running out of gas in the late game."),
    ("normal", "\u2022 Removal density (fraction of slots allocated to removal) \u2014 Decks need answers to threats; too few means vulnerability, too many means no proactive game plan."),
    ("normal", "\u2022 Card advantage density (fraction of slots allocated to draw/advantage) \u2014 The primary resource axis in Commander; running out of cards means losing."),
    ("normal", "\u2022 Synergy concentration (picks aligned with commander archetype) \u2014 Focused decks outperform unfocused \u201cgood stuff\u201d piles."),
    ("normal", "\u2022 Ramp density (fraction of slots for mana acceleration) \u2014 Commander games are won by deploying threats faster than opponents."),
    ("normal", "\u2022 Slot progress (early / mid / late in the episode) \u2014 What the deck needs changes as it fills; ramp is valuable early, finishers late."),
    ("normal", ""),
    ("normal", "With 6 variables \u00d7 3 bins each, the total state space is 3^6 = 729 possible states. This is deliberately coarse: fine enough to distinguish meaningfully different deck compositions, and small enough for tabular convergence."),
    ("normal", ""),
    ("heading3", "2.4 Action Space"),
    ("normal", "Each action corresponds to a card category \u2014 a functional abstraction over individual cards. Approximately 50 categories cover the roles needed in a Commander deck:"),
    ("normal", ""),
    ("normal", "\u2022 Lands: land_fetch, land_shock, land_dual, land_triome, land_utility, land_mdfc, land_basic"),
    ("normal", "\u2022 Mana: ramp_creature, ramp_spell, mana_rock, mana_dork"),
    ("normal", "\u2022 Removal: spot_removal_creature, spot_removal_spell, board_wipe, artifact_removal, enchantment_removal"),
    ("normal", "\u2022 Card advantage: card_draw_spell, card_draw_engine, impulse_draw, tutor, recursion"),
    ("normal", "\u2022 Counterspells: counterspell_hard, counterspell_soft"),
    ("normal", "\u2022 Creatures: creature_aggro, creature_utility, creature_evasive, creature_token_producer"),
    ("normal", "\u2022 Protection: protection_self, protection_permanent, hexproof_source"),
    ("normal", "\u2022 Synergy: tribal_payoff, tribal_lord, combo_piece, sacrifice_outlet, aristocrat_payoff"),
    ("normal", "\u2022 Value engines: enchantment_value, artifact_value, planeswalker"),
    ("normal", "\u2022 Finishers: wincon_combat, wincon_combo, wincon_alternate"),
    ("normal", ""),
    ("normal", "When the agent selects a category, a card selection heuristic chooses the best available card from the collection within that category, considering synergy with the commander, standalone card power, and what is already in the deck. This two-tier design (RL selects categories, heuristic selects cards) reduces the action space from thousands of individual cards to approximately 50 categories, making the problem tractable for tabular methods."),
    ("normal", ""),
    ("heading3", "2.5 Reward Function"),
    ("normal", "The reward function combines intermediate shaping rewards during the episode with a terminal reward at the end."),
    ("normal", ""),
    ("bold", "Intermediate rewards (per step):"),
    ("normal", "\u2022 Mana curve bonus/penalty: Small positive reward for picks that keep average CMC in a healthy range (2.5\u20133.5); small penalty for pushing it too high or too low."),
    ("normal", "\u2022 Diversity bonus: Reward for picking a category that addresses an underrepresented deck need."),
    ("normal", "\u2022 Synergy bonus: Reward for picks that align with the commander\u2019s archetype profile."),
    ("normal", "\u2022 Redundancy penalty: Penalty for over-allocating to a category that is already well-represented."),
    ("normal", ""),
    ("normal", "Intermediate shaping rewards are kept small relative to the terminal reward to avoid overwhelming the true objective."),
    ("normal", ""),
    ("bold", "Terminal reward (end of episode):"),
    ("normal", "Computed by submitting the completed decklist to the /synergy/analyze-deck API endpoint, which returns a composite score incorporating synergy rating, power heuristics (fast mana density, interaction density, card advantage density, average CMC, combo potential), and role balance classification. This terminal signal evaluates the complete deck holistically and serves as the primary learning signal."),
    ("normal", ""),
    ("heading3", "2.6 Transition Dynamics"),
    ("normal", "Transitions are deterministic given the agent\u2019s action and the heuristic\u2019s card selection:"),
    ("normal", "1. Agent observes state s_t (binned deck composition)."),
    ("normal", "2. Agent selects action a_t (card category)."),
    ("normal", "3. Heuristic selects the best available card from that category in the collection."),
    ("normal", "4. Card is added to the deck; environment computes new state s_{t+1}."),
    ("normal", "5. Intermediate reward r_t is computed."),
    ("normal", "6. If all meaningful slots are filled, the episode ends and the terminal reward is computed."),
    ("normal", ""),
    ("normal", "Because the environment is largely deterministic, the Dyna-Q model\u2019s predictions are highly accurate, and planning steps are very effective at accelerating learning."),
    ("normal", ""),
    ("heading3", "2.7 Scalability Justification"),
    ("normal", ""),
    ("bold", "Table 2: State-Action Space Comparison"),
    ("normal", "\u2022 Individual cards with continuous state: ~10^24 states \u00d7 3,000\u20138,000 actions = intractable"),
    ("normal", "\u2022 Individual cards with binned state: ~729 states \u00d7 3,000\u20138,000 actions = ~5.8M Q-table entries (borderline)"),
    ("normal", "\u2022 Card categories with binned state (this project): ~729 states \u00d7 ~50 actions = ~36,450 Q-table entries (tractable)"),
    ("normal", ""),
    ("normal", "The category abstraction reduces the problem by approximately five orders of magnitude while preserving the strategic decisions that determine deck quality."),
    ("normal", ""),
    ("heading2", "References"),
    ("normal", ""),
    ("normal", "Sutton, R. S. (1991). Dyna, an integrated architecture for learning, planning, and reacting. ACM SIGART Bulletin, 2(4), 160\u2013163."),
    ("normal", ""),
    ("normal", "Sutton, R. S., & Barto, A. G. (2018). Reinforcement learning: An introduction (2nd ed.). MIT Press."),
    ("normal", ""),
    ("normal", "Towers, M., Kwiatkowski, A., Terry, J., Balis, J. U., De Cola, G., Deleu, T., et al. (2024). Gymnasium: A standard interface for reinforcement learning environments. arXiv preprint arXiv:2407.17032."),
    ("normal", ""),
    ("normal", "Wizards of the Coast. (2024). Magic: The Gathering comprehensive rules \u2014 Commander format. https://magic.wizards.com/en/rules"),
    ("normal", ""),
]

# Insert content in reverse order after pa1_note
anchor = pa1_note
# We need to insert in reverse since addnext puts things right after the anchor
for block_type, text in reversed(content_blocks):
    if block_type == "heading2":
        p = doc.add_paragraph()
        p.style = doc.styles['Heading 2']
        p.text = text
        anchor._element.addnext(p._element)
    elif block_type == "heading3":
        p = doc.add_paragraph()
        p.style = doc.styles['Heading 3']
        p.text = text
        anchor._element.addnext(p._element)
    elif block_type == "bold":
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        anchor._element.addnext(p._element)
    else:
        p = doc.add_paragraph(text)
        anchor._element.addnext(p._element)

output_path = "deliverables/Reinforcement_Learning_Project_Week1.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
